"""Load corpus files with DirectoryLoader + TextLoader / PyPDFLoader."""

from __future__ import annotations

from pathlib import Path
from typing import Dict, Iterable, List

from langchain_community.document_loaders import DirectoryLoader, PyPDFLoader, TextLoader
from langchain_core.documents import Document

from RAG.configuration.rag_config_loader import DocumentsConfig, documents_config

_SKIP_NAMES = {"readme.md", "readme.txt", ".gitkeep"}


def _load_with(directory: Path, globs: tuple[str, ...], loader_cls, loader_kwargs=None) -> List[Document]:
    if not directory.exists():
        return []
    docs: List[Document] = []
    for pattern in globs:
        loader = DirectoryLoader(
            path=str(directory),
            glob=pattern,
            loader_cls=loader_cls,
            loader_kwargs=loader_kwargs or {},
            show_progress=False,
            use_multithreading=False,
            silent_errors=True,
        )
        docs.extend(loader.load())
    return docs


def _matching_files(directory: Path, globs: Iterable[str]) -> List[Path]:
    if not directory.exists():
        return []
    return sorted({path for pattern in globs for path in directory.glob(pattern) if path.is_file()})


def _load_docx(directory: Path, globs: tuple[str, ...]) -> List[Document]:
    """DOCX metnini doğrudan okur; eski DOC ve görüntü dosyaları bilinçli atlanır."""
    try:
        from docx import Document as WordDocument
    except ImportError:
        return []
    documents: List[Document] = []
    for path in _matching_files(directory, globs):
        try:
            word = WordDocument(path)
        except Exception:
            # Bozuk/şifreli dosya corpus'u durdurmaz; OCR veya dönüştürme yapılmaz.
            continue
        parts = [paragraph.text.strip() for paragraph in word.paragraphs if paragraph.text.strip()]
        for table in word.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        text = "\n".join(parts).strip()
        if text:
            documents.append(Document(page_content=text, metadata={"source": str(path), "extraction_method": "docx"}))
    return documents


def _load_xlsx(directory: Path, globs: tuple[str, ...]) -> List[Document]:
    """XLSX sayfalarını satır yapısını koruyan metin olarak indeksler."""
    try:
        from openpyxl import load_workbook
    except ImportError:
        return []
    documents: List[Document] = []
    for path in _matching_files(directory, globs):
        try:
            book = load_workbook(path, read_only=True, data_only=True)
        except Exception:
            continue
        try:
            pages: List[str] = []
            for sheet in book.worksheets:
                rows = [" | ".join(str(value).strip() for value in row if value not in (None, "")) for row in sheet.iter_rows(values_only=True)]
                rows = [row for row in rows if row]
                if rows:
                    pages.append(f"[SAYFA: {sheet.title}]\n" + "\n".join(rows))
            text = "\n\n".join(pages).strip()
            if text:
                documents.append(Document(page_content=text, metadata={"source": str(path), "extraction_method": "xlsx"}))
        finally:
            book.close()
    return documents


def _bucket_name(directory: Path, cfg: DocumentsConfig) -> str:
    resolved = directory.resolve()
    mapping = {
        cfg.laws_path.resolve(): "laws",
        cfg.regulations_path.resolve(): "regulations",
        cfg.amendments_path.resolve(): "amendments",
        cfg.internal_docs_path.resolve(): "internal_docs",
        cfg.uploads_path.resolve(): "uploads",
        cfg.reference_docs_path.resolve(): "reference_docs",
        cfg.classification_data_path.resolve(): "reference_docs",
    }
    return mapping.get(resolved, "unknown")


def _keep(path: Path) -> bool:
    name = path.name.lower()
    return name not in _SKIP_NAMES and not name.startswith(".") and not name.endswith(".meta.json")


def _document_category(source: Path, directory: Path, source_type: str, cfg: DocumentsConfig) -> str:
    """Klasör etiketi varsa kullanır; kanun kaynaklarında güvenli varsayılan döner."""
    if directory.resolve() == cfg.classification_data_path.resolve():
        try:
            return source.resolve().relative_to(cfg.classification_data_path.resolve()).parts[0]
        except (ValueError, IndexError):
            return "unknown"
    defaults = {
        "laws": "law", "regulations": "regulation", "amendments": "amendment",
        "reference_docs": "reference_document", "internal_docs": "internal_document", "uploads": "uploaded_document",
    }
    return defaults.get(source_type, "unknown")


def _merge_pdf_pages(documents: List[Document]) -> List[Document]:
    """Merge PDF pages before legal article extraction.

    ``PyPDFLoader`` yields one Document per page.  Treating each page as an
    independent legal document turns article continuations into false
    preambles, loses the article number, and produces unstable chunk IDs.
    Text documents pass through unchanged; PDF pages are merged in page order
    and retain their original page range for citations.
    """
    grouped: Dict[str, List[Document]] = {}
    passthrough: List[Document] = []
    for doc in documents:
        source = Path(str((doc.metadata or {}).get("source", "")))
        if source.suffix.lower() != ".pdf":
            passthrough.append(doc)
            continue
        grouped.setdefault(str(source), []).append(doc)

    merged: List[Document] = list(passthrough)
    for source, pages in grouped.items():
        pages.sort(key=lambda page: int((page.metadata or {}).get("page", 0)))
        first_meta = dict(pages[0].metadata or {})
        page_numbers = [int((page.metadata or {}).get("page", 0)) for page in pages]
        first_meta.update(
            {
                "source": source,
                "page_start": min(page_numbers) + 1,
                "page_end": max(page_numbers) + 1,
            }
        )
        merged.append(
            Document(
                # Kanun tek belge olarak işlenirken sayfa sınırları korunur.
                # Hukukî chunker özel işaretleri tüketir ve her chunk'a sayfa aralığı yazar.
                page_content="\n\n".join(
                    f"[[RAG_PAGE:{int((page.metadata or {}).get('page', 0)) + 1}]]\n{page.page_content}"
                    for page in pages if page.page_content
                ),
                metadata=first_meta,
            )
        )
    return merged


def load_directory(directory: Path, cfg: DocumentsConfig | None = None) -> List[Document]:
    cfg = cfg or documents_config
    raw = _load_with(directory, cfg.text_globs, TextLoader, {"encoding": "utf-8"})
    raw += _load_with(directory, cfg.pdf_globs, PyPDFLoader)
    raw += _load_docx(directory, cfg.docx_globs)
    raw += _load_xlsx(directory, cfg.spreadsheet_globs)

    raw = _merge_pdf_pages(raw)
    source_type = _bucket_name(directory, cfg)
    kept: List[Document] = []
    for doc in raw:
        meta = dict(doc.metadata or {})
        source = Path(str(meta.get("source", "")))
        if not _keep(source):
            continue
        if not doc.page_content.strip():
            continue
        meta.update(
            {
                "source_type": source_type,
                "source_dir": str(directory),
                "source_file": source.name or "unknown",
                "law_name": meta.get("law_name") or source.stem or "unknown",
                "document_category": meta.get("document_category") or _document_category(source, directory, source_type, cfg),
            }
        )
        doc.metadata = meta
        kept.append(doc)
    return kept


def load_all_sources(cfg: DocumentsConfig | None = None) -> List[Document]:
    cfg = cfg or documents_config
    docs: List[Document] = []
    for directory in cfg.all_source_dirs:
        docs.extend(load_directory(directory, cfg))
    return docs
